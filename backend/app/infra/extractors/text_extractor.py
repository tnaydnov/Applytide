"""
Text Extraction Module for Multiple Document Formats

This module provides unified text extraction for various document formats:
- Plain text (.txt)
- PDF (.pdf) - via PDFExtractorPort
- Word documents (.docx)
- Legacy Word documents (.doc) - notice only

Features:
    - Format detection based on file extension
    - Automatic encoding handling for text files
    - Table extraction from Word documents
    - Bullet point normalization
    - Comprehensive error handling
    - Detailed logging

Supported Formats:
    - .txt: UTF-8 with fallback encoding detection
    - .pdf: Delegated to PDFExtractorPort implementation
    - .docx: Full support (text + tables)
    - .doc: Legacy format (notice returned, manual conversion required)

Error Handling:
    - File not found → descriptive error message
    - Encoding errors → fallback to 'ignore' mode
    - Corrupted files → error message with details
    - Missing dependencies → clear error message

Author: ApplyTide Team
Last Updated: 2025-01-18
"""
from __future__ import annotations
from pathlib import Path
import re
from typing import Protocol, Optional

from ..logging import get_logger

logger = get_logger(__name__)

# Constants
SAFE_BULLET = "•"  # Normalized bullet character
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10MB max for text files
ENCODING_FALLBACKS = ["utf-8", "latin-1", "cp1252", "utf-16"]


class TextExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


class PDFExtractorPort(Protocol):
    """Port interface for PDF extraction."""
    def extract_text(self, pdf_content: bytes) -> dict: ...

class TextExtractor:
    """
    Unified text extraction for multiple document formats.
    
    Handles text extraction from various file formats with automatic
    format detection, encoding handling, and error recovery.
    
    Supported Formats:
        - .txt: Plain text with automatic encoding detection
        - .pdf: Delegated to PDFExtractorPort (injected dependency)
        - .docx: Microsoft Word (modern format)
        - .doc: Legacy Word (notice only, conversion required)
    
    Features:
        - Automatic format detection based on extension
        - Encoding fallback for text files
        - Table extraction from Word documents
        - Bullet point normalization
        - Comprehensive error handling
    
    Args:
        pdf_extractor: PDFExtractorPort implementation for PDF handling
    
    Thread Safety:
        - Thread-safe (no shared mutable state)
        - Safe for concurrent use
    
    Example:
        >>> from pdf_extractor import PDFExtractor
        >>> pdf_ext = PDFExtractor()
        >>> text_ext = TextExtractor(pdf_extractor=pdf_ext)
        >>> text = text_ext.extract_text(Path("resume.pdf"))
    """
    
    def __init__(self, pdf_extractor: PDFExtractorPort) -> None:
        """
        Initialize TextExtractor with PDF extraction capability.
        
        Args:
            pdf_extractor: PDFExtractorPort implementation for handling PDF files
        
        Raises:
            TypeError: If pdf_extractor doesn't implement PDFExtractorPort
        """
        if not pdf_extractor:
            raise TypeError("pdf_extractor is required")
        
        self.pdf_extractor = pdf_extractor
        
        logger.debug("TextExtractor initialized")

    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from document file.
        
        Automatically detects format based on file extension and applies
        appropriate extraction method. Returns extracted text or error message.
        
        Args:
            file_path: Path to document file
        
        Returns:
            str: Extracted text content, or error message if extraction fails
        
        Supported Extensions:
            - .txt: Plain text (UTF-8, Latin-1, CP1252, UTF-16 fallbacks)
            - .pdf: PDF documents (via PDFExtractorPort)
            - .docx: Modern Word documents (text + tables)
            - .doc: Legacy Word documents (returns conversion notice)
        
        Notes:
            - Returns error messages as strings (doesn't raise exceptions)
            - Empty files return empty string
            - Unknown extensions return empty string
            - File must exist and be readable
        
        Example:
            >>> extractor = TextExtractor(pdf_extractor)
            >>> text = extractor.extract_text(Path("resume.pdf"))
            >>> if text.startswith("Error"):
            ...     print("Extraction failed")
        """
        # Validate input
        if not file_path:
            logger.error("extract_text called with None file_path")
            return "Error extracting text: file_path is None"
        
        if not isinstance(file_path, Path):
            logger.error(
                "Invalid file_path type",
                extra={"type": type(file_path).__name__}
            )
            return f"Error extracting text: file_path must be Path, got {type(file_path).__name__}"
        
        # Check file exists
        if not file_path.exists():
            logger.error(
                "File not found",
                extra={"path": str(file_path)}
            )
            return f"Error extracting text: File not found: {file_path}"
        
        # Check file is readable
        if not file_path.is_file():
            logger.error(
                "Path is not a file",
                extra={"path": str(file_path)}
            )
            return f"Error extracting text: Path is not a file: {file_path}"
        
        # Get file extension
        ext = file_path.suffix.lower()
        
        if not ext:
            logger.warning(
                "File has no extension",
                extra={"path": str(file_path)}
            )
            return "Error extracting text: File has no extension"
        
        logger.debug(
            "Starting text extraction",
            extra={
                "path": str(file_path),
                "extension": ext,
                "size_bytes": file_path.stat().st_size
            }
        )
        
        try:
            # Extract based on format
            if ext == ".txt":
                return self._extract_txt(file_path)
            
            if ext == ".pdf":
                return self._extract_pdf(file_path)
            
            if ext in (".docx", ".doc"):
                return self._extract_word(file_path, ext)
            
            # Unknown extension
            logger.warning(
                "Unsupported file extension",
                extra={"extension": ext, "path": str(file_path)}
            )
            return f"Error extracting text: Unsupported file format: {ext}"
        
        except Exception as e:
            logger.error(
                "Unexpected error during text extraction",
                extra={
                    "path": str(file_path),
                    "extension": ext,
                    "error": str(e),
                    "error_type": type(e).__name__
                },
                exc_info=True
            )
            return f"Error extracting text: {e}"
    
    def _extract_txt(self, file_path: Path) -> str:
        """
        Extract text from .txt file with encoding detection.
        
        Tries multiple encodings in order: UTF-8, Latin-1, CP1252, UTF-16.
        Falls back to 'ignore' mode if all encodings fail.
        
        Args:
            file_path: Path to .txt file
        
        Returns:
            str: Extracted text content
        """
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > MAX_FILE_SIZE_BYTES:
            logger.warning(
                "Text file too large",
                extra={
                    "size_bytes": file_size,
                    "max_bytes": MAX_FILE_SIZE_BYTES,
                    "path": str(file_path)
                }
            )
            return f"Error extracting text: File too large ({file_size / (1024*1024):.1f}MB). Maximum: {MAX_FILE_SIZE_BYTES / (1024*1024):.0f}MB"
        
        # Try encodings in order
        for encoding in ENCODING_FALLBACKS:
            try:
                text = file_path.read_text(encoding=encoding, errors="strict")
                logger.debug(
                    f"Successfully read text file with {encoding}",
                    extra={
                        "encoding": encoding,
                        "char_count": len(text)
                    }
                )
                return text
            
            except (UnicodeDecodeError, LookupError):
                continue
        
        # All encodings failed, use ignore mode
        logger.warning(
            "All encodings failed, using ignore mode",
            extra={"path": str(file_path)}
        )
        try:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            logger.error(
                "Failed to read text file even with ignore mode",
                extra={"path": str(file_path), "error": str(e)}
            )
            return f"Error extracting text: Unable to read file: {e}"
    
    def _extract_pdf(self, file_path: Path) -> str:
        """
        Extract text from .pdf file via PDFExtractorPort.
        
        Args:
            file_path: Path to .pdf file
        
        Returns:
            str: Extracted text content or error message
        """
        try:
            data = file_path.read_bytes()
            logger.debug(
                "Read PDF bytes",
                extra={"size_bytes": len(data)}
            )
            
            res = self.pdf_extractor.extract_text(data)
            
            if not isinstance(res, dict):
                logger.error(
                    "PDF extractor returned invalid type",
                    extra={"type": type(res).__name__}
                )
                return "Error extracting text: PDF extractor returned invalid result"
            
            text = res.get("text", "")
            success = res.get("success", False)
            
            if not success:
                error = res.get("error", "Unknown error")
                logger.warning(
                    "PDF extraction failed",
                    extra={"error": error}
                )
                return f"Error extracting text from PDF: {error}"
            
            logger.debug(
                "PDF extraction successful",
                extra={
                    "char_count": len(text),
                    "method": res.get("extraction_method", "unknown")
                }
            )
            
            return text
        
        except Exception as e:
            logger.error(
                "Error reading or extracting PDF",
                extra={
                    "path": str(file_path),
                    "error": str(e),
                    "error_type": type(e).__name__
                },
                exc_info=True
            )
            return f"Error extracting text from PDF: {e}"
    
    def _extract_word(self, file_path: Path, ext: str) -> str:
        """
        Extract text from Word documents (.docx or .doc).
        
        Args:
            file_path: Path to Word document
            ext: File extension (".docx" or ".doc")
        
        Returns:
            str: Extracted text content or error/notice message
        """
        try:
            from docx import Document as _Docx
        except ImportError:
            logger.error("python-docx library not available")
            return "Error extracting text: python-docx library not installed"
        
        # Handle legacy .doc format
        if ext == ".doc":
            logger.info(
                "Legacy .doc format detected",
                extra={"path": str(file_path)}
            )
            return "Legacy .doc file: text extraction requires conversion to .docx"
        
        # Extract from .docx
        try:
            d = _Docx(str(file_path))
            parts: list[str] = []
            
            # Extract paragraphs
            for p in d.paragraphs:
                txt = p.text.strip()
                if txt:
                    parts.append(txt)
            
            # Extract tables
            for t in d.tables:
                for row in t.rows:
                    row_txt = " | ".join(
                        cell.text.strip() 
                        for cell in row.cells 
                        if cell and cell.text
                    )
                    if row_txt:
                        parts.append(row_txt)
            
            text = "\n".join(parts)
            
            # Normalize bullet points
            text = re.sub(r"^[\u2022\u25CF\-]\s*", SAFE_BULLET + " ", text, flags=re.M)
            
            logger.debug(
                "Word document extraction successful",
                extra={
                    "paragraph_count": len(d.paragraphs),
                    "table_count": len(d.tables),
                    "char_count": len(text)
                }
            )
            
            return text
        
        except Exception as e:
            logger.error(
                "Error extracting from Word document",
                extra={
                    "path": str(file_path),
                    "error": str(e),
                    "error_type": type(e).__name__
                },
                exc_info=True
            )
            
            # Check if it's actually a .doc file with wrong extension
            if "not a zip file" in str(e).lower() or "bad zip" in str(e).lower():
                return "Legacy .doc file detected: text extraction requires conversion to .docx"
            
            return f"Error extracting text from Word document: {e}"
